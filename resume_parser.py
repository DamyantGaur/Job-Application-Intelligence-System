"""
resume_parser.py — PDF and DOCX Resume Text Extraction

Provides a unified interface for extracting clean plain text from
uploaded resume files in PDF or DOCX format.
"""

import os
import re
from io import BytesIO

from pdfminer.high_level import extract_text as pdf_extract_text
from docx import Document


def parse_resume(file) -> str:
    """
    Extract plain text from a resume file (PDF or DOCX).

    Args:
        file: A file-like object (e.g., Streamlit UploadedFile) or a file path string.
              The file name must end with .pdf or .docx.

    Returns:
        A cleaned plain-text string with normalised whitespace.

    Raises:
        ValueError: If the file type is not PDF or DOCX.
    """
    # Determine the file name for extension detection
    if isinstance(file, str):
        file_name = file
    else:
        file_name = getattr(file, "name", "")

    ext = os.path.splitext(file_name)[1].lower()

    if ext == ".pdf":
        raw_text = _extract_pdf(file)
    elif ext == ".docx":
        raw_text = _extract_docx(file)
    elif ext == ".txt":
        raw_text = _extract_txt(file)
    else:
        raise ValueError(
            f"Unsupported file type '{ext}'. Please upload a PDF, DOCX, or TXT file."
        )

    return _clean_text(raw_text)


def _extract_pdf(file) -> str:
    """
    Extract text from a PDF file using pdfminer.six.

    Args:
        file: A file path string or a file-like object with a .read() method.

    Returns:
        Raw extracted text from all pages of the PDF.
    """
    if isinstance(file, str):
        return pdf_extract_text(file)
    else:
        # Streamlit UploadedFile — read bytes and wrap in BytesIO
        content = file.read()
        file.seek(0)  # Reset for potential re-reads
        return pdf_extract_text(BytesIO(content))


def _extract_docx(file) -> str:
    """
    Extract text from a DOCX file using python-docx.

    Args:
        file: A file path string or a file-like object.

    Returns:
        Raw extracted text with paragraphs joined by newlines.
    """
    if isinstance(file, str):
        doc = Document(file)
    else:
        content = file.read()
        file.seek(0)
        doc = Document(BytesIO(content))

    paragraphs = [para.text for para in doc.paragraphs if para.text.strip()]
    return "\n".join(paragraphs)


def _extract_txt(file) -> str:
    """
    Extract text from a plain-text (.txt) file.

    Uses UTF-8 decoding with error replacement to handle corrupt characters.

    Args:
        file: A file path string or a file-like object.

    Returns:
        Decoded text content.
    """
    if isinstance(file, str):
        with open(file, "r", encoding="utf-8", errors="replace") as f:
            return f.read()
    else:
        content = file.read()
        file.seek(0)
        if isinstance(content, bytes):
            return content.decode("utf-8", errors="replace")
        return content


def _clean_text(text: str) -> str:
    """
    Normalise whitespace and remove non-printable characters.

    Args:
        text: Raw extracted text.

    Returns:
        Cleaned text with collapsed whitespace and trimmed lines.
    """
    # Sanitize: force through UTF-8 to strip hidden ligatures / corrupt chars
    text = text.encode("utf-8", errors="ignore").decode("utf-8")
    # Remove non-printable characters (keep newlines, tabs, spaces)
    text = re.sub(r"[^\S\n]+", " ", text)
    # Collapse multiple blank lines into one
    text = re.sub(r"\n{3,}", "\n\n", text)
    # Strip each line
    lines = [line.strip() for line in text.splitlines()]
    return "\n".join(lines).strip()
